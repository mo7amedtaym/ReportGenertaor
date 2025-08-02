import os
from datetime import datetime
import tkinter as tk
from tkinter import messagebox
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import sys

# مسار القالب
TEMPLATE_PATH = "temp1.docx"
OUTPUT_FOLDER = "قضايا"

# تأكد من وجود مجلد الحفظ
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# المتغيرات المستخدمة داخل القالب
PLACEHOLDERS = {
    "مدير النيابة":"{{pro}}",
    "سكرتير التحقيق":"{{sec}}",
    "اسم المتهم": "{{name}}",
    "رقم القضية": "{{num}}",
    "(السنه)تاريخ القضية": "{{y}}",
    "الشهر": "{{m}}",
    "اليوم": "{{d}}",
    "الساعة": "{{h}}",
    "الدقيقة": "{{min}}",
    "رقم المركبة": "{{plate}}",
    "نوع المركبة": "{{car_type}}",
    "عنوان المتهم": "{{address}}",
    "عمر المتهم": "{{age}}",
    "الوظيفة او المؤهل":"{{job}}",
    "الرقم القومي": "{{id}}",
    "اسم الضابط": "{{cap_name}}",
    "رتبته": "{{cap_title}}",
    "القسم المسؤول": "{{department}}",
}

def generate_document(values):
    doc = Document(TEMPLATE_PATH)
    for p in doc.paragraphs:
        for key, placeholder in PLACEHOLDERS.items():
            if placeholder in p.text:
                p.text = p.text.replace(placeholder, values[key])

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, placeholder in PLACEHOLDERS.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, values[key])

    # ضبط الخط لجميع الفقرات بعد التعديل
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = 'Simplified Arabic'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Simplified Arabic')
            run.font.size = Pt(18)

    # ضبط الخط للنص داخل الجداول إن وجدت
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = 'Simplified Arabic'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Simplified Arabic')
                        run.font.size = Pt(18)
    filename = f"{values['اسم المتهم'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    output_path = os.path.join(OUTPUT_FOLDER, filename)
    doc.save(output_path)
    return output_path

def create_gui(output_filename, suspect_name, other_details):
    window = tk.Tk()
    window.title(output_filename)
    window.geometry("550x650")

    entries = {}
    row = 0
    for key in PLACEHOLDERS.keys():
        label = tk.Label(window, text=key + ":", anchor="w", font=("Arial", 12))
        label.grid(row=row, column=0, sticky="w", padx=10, pady=5)
        entry = tk.Entry(window, width=40, font=("Arial", 12))
        entry.grid(row=row, column=1, padx=10, pady=5)
        entries[key] = entry
        row += 1

    def on_generate():
        values = {key: entry.get().strip() for key, entry in entries.items()}
        if not all(values.values()):
            messagebox.showerror("خطأ", other_details)
            return
        output_path = generate_document(values)
        messagebox.showinfo("تم", f"تم إنشاء المستند:\n{output_path}")
        os.startfile(output_path)

    generate_btn = tk.Button(window, text="توليد المستند وفتحه", command=on_generate, bg="#4CAF50", fg="white", font=("Arial", 14))
    generate_btn.grid(row=row, column=0, columnspan=2, pady=20)

    window.mainloop()

if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("Usage: python create_police_report.py <output_filename> <suspect_name> <other_details>", file=sys.stderr)
        sys.exit(1)

    output_filename = sys.argv[1]
    suspect_name = sys.argv[2]
    other_details = sys.argv[3]

    try:
        create_gui(output_filename, suspect_name, other_details)
    except Exception as e:
        print(f"Error creating DOCX: {e}", file=sys.stderr)
        sys.exit(1)